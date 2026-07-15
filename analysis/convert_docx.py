"""Convert HumanEval/{FS,ABLMC}/InjTest_*.docx into plain-text card files.

Each output .txt has a small header (condition, metric, source filename),
the test-results table, and the participant's metric-card content (everything
after the first "Metric Card:" marker that doesn't also contain "Injection
Test", with trailing timestamp/duration lines stripped).
"""
import argparse
import re
from pathlib import Path

import docx


def find_card_start(paragraphs):
    for i, p in enumerate(paragraphs):
        t = p.text.strip()
        if "Metric Card:" in t and "Injection Test" not in t:
            return i + 1
    return None


def is_timestamp_line(s):
    s2 = s.strip()
    if not s2 or len(s2) > 60:
        return False
    if not re.search(r"\d", s2):
        return False
    return bool(
        ":" in s2 or "(" in s2 or
        re.search(r"\b(min|sek|uhr|millisek)", s2, re.I)
    )


QUESTIONNAIRE_RE = re.compile(r"^\s*(question+aire|questionnaire)\s*:?\s*$", re.I)


def extract_participant_card(doc):
    paras = doc.paragraphs
    start = find_card_start(paras)
    if start is None:
        return ""
    lines = [p.text.rstrip() for p in paras[start:]]
    for i, line in enumerate(lines):
        if QUESTIONNAIRE_RE.match(line):
            lines = lines[:i]
            break
    while lines and (not lines[-1].strip() or is_timestamp_line(lines[-1])):
        lines.pop()
    return "\n".join(lines).strip()


def extract_table(doc):
    out = []
    for t in doc.tables:
        for row in t.rows:
            out.append("\t".join(
                c.text.replace("\n", " ").strip() for c in row.cells
            ))
    return "\n".join(out)


def process(src, out_dir, condition, metric):
    doc = docx.Document(src)
    body = (
        f"# Condition: {condition}\n"
        f"# Metric: {metric}\n"
        f"# Source: {src.name}\n"
        f"\n## Test results table\n{extract_table(doc)}\n"
        f"\n## Participant metric card\n{extract_participant_card(doc)}\n"
    )
    out_path = out_dir / f"{src.stem}.txt"
    out_path.write_text(body, encoding="utf-8")
    return out_path


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--src", default="HumanEval")
    ap.add_argument("--out-dir", default="analysis/cards_text")
    args = ap.parse_args()

    src_dir = Path(args.src)
    out_dir = Path(args.out_dir); out_dir.mkdir(parents=True, exist_ok=True)

    for sub, condition, metric in [
        ("FS", "FS", "ROUGE-L"),
        ("ABLMC", "ABLMC", "BARTScore-avg-f"),
    ]:
        for f in sorted((src_dir / sub).glob("InjTest_*.docx")):
            print(f"  -> {process(f, out_dir, condition, metric)}")


if __name__ == "__main__":
    main()
